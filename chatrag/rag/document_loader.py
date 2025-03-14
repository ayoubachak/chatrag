from pathlib import Path
import uuid
import pandas as pd
from typing import List, Dict, Any, Literal, Tuple
from abc import ABC, abstractmethod
import re
import numpy as np
from collections import Counter


class ChunkingStrategy(ABC):
    """
    Abstract base class for document chunking strategies.
    """
    
    @abstractmethod
    async def chunk_text(self, text: str, metadata: Dict[str, Any]) -> List[Dict[str, str]]:
        """
        Chunk a text document into smaller pieces.
        
        Args:
            text: The text to chunk
            metadata: Metadata about the source document
            
        Returns:
            List of document chunks with metadata
        """
        pass
        
    @abstractmethod
    async def chunk_pdf(self, pdf_reader, file_path: str) -> List[Dict[str, str]]:
        """
        Chunk a PDF document into smaller pieces.
        
        Args:
            pdf_reader: The PDF reader object
            file_path: Path to the PDF file
            
        Returns:
            List of document chunks with metadata
        """
        pass
        
    @abstractmethod
    async def chunk_docx(self, doc, file_path: str) -> List[Dict[str, str]]:
        """
        Chunk a Word document into smaller pieces.
        
        Args:
            doc: The Word document object
            file_path: Path to the Word file
            
        Returns:
            List of document chunks with metadata
        """
        pass
        
    @abstractmethod
    async def chunk_tabular(self, df, file_path: str) -> List[Dict[str, str]]:
        """
        Chunk tabular data into smaller pieces.
        
        Args:
            df: The pandas DataFrame
            file_path: Path to the tabular file
            
        Returns:
            List of document chunks with metadata
        """
        pass

class BasicChunker(ChunkingStrategy):
    """
    Basic chunking strategy that splits text by paragraphs and processes
    documents page by page.
    """
    
    async def chunk_text(self, text: str, metadata: Dict[str, Any]) -> List[Dict[str, str]]:
        """
        Chunk a text document by paragraphs.
        """
        # Simple chunking by paragraphs
        chunks = [p.strip() for p in text.split("\n\n") if p.strip()]
        
        return [
            {
                "content": chunk,
                "metadata": {
                    **metadata,
                    "chunk_id": i
                }
            }
            for i, chunk in enumerate(chunks)
        ]
        
    async def chunk_pdf(self, pdf_reader, file_path: str) -> List[Dict[str, str]]:
        """
        Chunk a PDF document page by page.
        """
        text_chunks = []
        
        for i, page in enumerate(pdf_reader.pages):
            text = page.extract_text()
            if text.strip():
                text_chunks.append({
                    "content": text.strip(),
                    "metadata": {
                        "source": Path(file_path).name,
                        "page": i + 1,
                        "chunk_id": i
                    }
                })
                
        return text_chunks
        
    async def chunk_docx(self, doc, file_path: str) -> List[Dict[str, str]]:
        """
        Chunk a Word document by paragraphs.
        """
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        
        return [
            {
                "content": para,
                "metadata": {
                    "source": Path(file_path).name,
                    "chunk_id": i
                }
            }
            for i, para in enumerate(paragraphs)
        ]
        
    async def chunk_tabular(self, df, file_path: str) -> List[Dict[str, str]]:
        """
        Chunk tabular data row by row.
        """
        chunks = []
        for i, row in enumerate(df.to_dict("records")):
            # Format the row as a string
            content = "\n".join([f"{k}: {v}" for k, v in row.items()])
            chunks.append({
                "content": content,
                "metadata": {
                    "source": Path(file_path).name,
                    "row": i + 1,
                    "chunk_id": i
                }
            })
            
        return chunks

class SuperChunker(ChunkingStrategy):
    """
    Advanced chunking strategy that uses semantic boundaries and overlapping
    to create more meaningful chunks.
    """
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        """
        Initialize the super chunker.
        
        Args:
            chunk_size: Target size of each chunk in characters
            chunk_overlap: Overlap between chunks in characters
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        
    async def chunk_text(self, text: str, metadata: Dict[str, Any]) -> List[Dict[str, str]]:
        """
        Chunk a text document using semantic boundaries and overlapping.
        """
        # First split by paragraphs
        paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
        
        # Create chunks with target size and overlap
        chunks = []
        current_chunk = ""
        current_size = 0
        
        for i, para in enumerate(paragraphs):
            para_size = len(para)
            
            # If adding this paragraph would exceed the chunk size and we already have content,
            # save the current chunk and start a new one with overlap
            if current_size + para_size > self.chunk_size and current_chunk:
                chunks.append(current_chunk)
                
                # Find a good breaking point for the overlap
                # Start with the last self.chunk_overlap characters
                overlap_text = current_chunk[-self.chunk_overlap:] if len(current_chunk) > self.chunk_overlap else current_chunk
                
                # Try to find a sentence boundary in the overlap
                sentences = re.split(r'(?<=[.!?])\s+', overlap_text)
                if len(sentences) > 1:
                    # Start the new chunk with the last sentence(s) from the previous chunk
                    current_chunk = sentences[-1]
                else:
                    # If no sentence boundary, just use the overlap
                    current_chunk = overlap_text
                
                current_size = len(current_chunk)
            
            # Add the paragraph to the current chunk
            if current_chunk:
                current_chunk += "\n\n" + para
                current_size += para_size + 2  # +2 for the newlines
            else:
                current_chunk = para
                current_size = para_size
                
        # Add the last chunk if it's not empty
        if current_chunk:
            chunks.append(current_chunk)
            
        # Create the final chunk objects with metadata
        return [
            {
                "content": chunk,
                "metadata": {
                    **metadata,
                    "chunk_id": i,
                    "chunk_strategy": "semantic"
                }
            }
            for i, chunk in enumerate(chunks)
        ]
        
    async def chunk_pdf(self, pdf_reader, file_path: str) -> List[Dict[str, str]]:
        """
        Chunk a PDF document using semantic boundaries across pages.
        """
        # First extract text from all pages
        all_text = ""
        page_breaks = []
        
        for i, page in enumerate(pdf_reader.pages):
            page_text = page.extract_text()
            if page_text.strip():
                all_text += page_text + "\n\n"
                page_breaks.append(len(all_text))
                
        # Use the text chunking method on the combined text
        base_metadata = {"source": Path(file_path).name}
        chunks = await self.chunk_text(all_text, base_metadata)
        
        # Add page information to each chunk
        for chunk in chunks:
            # Find which page(s) this chunk belongs to
            chunk_start = all_text.find(chunk["content"])
            chunk_end = chunk_start + len(chunk["content"])
            
            # Find the page number(s)
            start_page = 1
            end_page = 1
            
            for i, break_pos in enumerate(page_breaks):
                if chunk_start < break_pos:
                    start_page = i + 1
                    break
                    
            for i, break_pos in enumerate(page_breaks):
                if chunk_end <= break_pos:
                    end_page = i + 1
                    break
            
            # Update metadata
            chunk["metadata"]["page_start"] = start_page
            chunk["metadata"]["page_end"] = end_page
            
        return chunks
        
    async def chunk_docx(self, doc, file_path: str) -> List[Dict[str, str]]:
        """
        Chunk a Word document using semantic boundaries.
        """
        # Extract all paragraphs
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        
        # Combine paragraphs into a single text
        all_text = "\n\n".join(paragraphs)
        
        # Use the text chunking method
        base_metadata = {"source": Path(file_path).name}
        return await self.chunk_text(all_text, base_metadata)
        
    async def chunk_tabular(self, df, file_path: str) -> List[Dict[str, str]]:
        """
        Chunk tabular data by groups of related rows.
        """
        chunks = []
        records = df.to_dict("records")
        
        # Group rows into chunks of appropriate size
        current_chunk = []
        current_size = 0
        
        for i, row in enumerate(records):
            # Format the row as a string
            row_text = "\n".join([f"{k}: {v}" for k, v in row.items()])
            row_size = len(row_text)
            
            # If adding this row would exceed the chunk size and we already have content,
            # save the current chunk and start a new one
            if current_size + row_size > self.chunk_size and current_chunk:
                # Combine the rows into a single chunk
                content = "\n\n".join(current_chunk)
                chunks.append({
                    "content": content,
                    "metadata": {
                        "source": Path(file_path).name,
                        "rows": f"{i-len(current_chunk)+1}-{i}",
                        "chunk_id": len(chunks),
                        "chunk_strategy": "semantic"
                    }
                })
                
                # Start a new chunk
                current_chunk = []
                current_size = 0
                
            # Add the row to the current chunk
            current_chunk.append(row_text)
            current_size += row_size
                
        # Add the last chunk if it's not empty
        if current_chunk:
            content = "\n\n".join(current_chunk)
            chunks.append({
                "content": content,
                "metadata": {
                    "source": Path(file_path).name,
                    "rows": f"{len(records)-len(current_chunk)+1}-{len(records)}",
                    "chunk_id": len(chunks),
                    "chunk_strategy": "semantic"
                }
            })
            
        return chunks
    
class QuantumChunker(ChunkingStrategy):
    """
    Highly advanced chunking strategy with intelligent boundary detection,
    dynamic overlapping, and content-aware processing for optimal RAG performance.
    
    Features:
    - Smart sentence and paragraph detection
    - Semantic coherence analysis
    - Header/section detection
    - Dynamic overlap based on content complexity
    - Context-aware metadata enrichment
    - Multi-modal content support (text, tables, code blocks)
    - Special handling for various document structures
    - Entity and concept detection to maintain context
    """
    
    def __init__(self, 
                chunk_size: int = 1000, 
                min_chunk_size: int = 200,
                max_chunk_size: int = 2000,
                chunk_overlap: int = 200,
                respect_sections: bool = True,
                respect_entities: bool = True,
                dynamic_overlap: bool = True):
        """
        Initialize the quantum chunker with advanced options.
        
        Args:
            chunk_size: Target size of each chunk in characters
            min_chunk_size: Minimum allowed chunk size
            max_chunk_size: Maximum allowed chunk size
            chunk_overlap: Base overlap between chunks in characters
            respect_sections: Whether to avoid breaking across section boundaries
            respect_entities: Whether to avoid breaking named entities
            dynamic_overlap: Whether to adjust overlap based on content complexity
        """
        self.chunk_size = chunk_size
        self.min_chunk_size = min_chunk_size
        self.max_chunk_size = max_chunk_size
        self.chunk_overlap = chunk_overlap
        self.respect_sections = respect_sections
        self.respect_entities = respect_entities
        self.dynamic_overlap = dynamic_overlap
        
        # Regex patterns for document structure detection
        self.section_pattern = re.compile(r'^#+\s+(.+)$|^(.+)\n[=\-]{3,}$', re.MULTILINE)
        self.list_pattern = re.compile(r'^\s*(\d+\.|[\*\-\+])\s+', re.MULTILINE)
        self.code_block_pattern = re.compile(r'```.*?```', re.DOTALL)
        self.table_pattern = re.compile(r'\|\s+.+\s+\|.*\n\|\s*[\-:]+\s*\|', re.MULTILINE)
        
        # Try to load NLP libraries if available
        self.nlp = None
        self.sentence_tokenizer = None
        try:
            import spacy
            self.nlp = spacy.load("en_core_web_sm")
        except (ImportError, OSError):
            try:
                import nltk
                nltk.download('punkt', quiet=True)
                from nltk.tokenize import sent_tokenize
                self.sentence_tokenizer = sent_tokenize
            except ImportError:
                # Fall back to regex-based sentence tokenization
                self.sentence_tokenizer = lambda text: re.split(r'(?<=[.!?])\s+', text)
        
    def _tokenize_sentences(self, text: str) -> List[str]:
        """
        Tokenize text into sentences using the best available method.
        
        Args:
            text: Text to tokenize
            
        Returns:
            List of sentences
        """
        if self.nlp:
            doc = self.nlp(text)
            return [str(sent) for sent in doc.sents]
        elif self.sentence_tokenizer:
            return self.sentence_tokenizer(text)
        else:
            # Fallback to regex
            return re.split(r'(?<=[.!?])\s+', text)
    
    def _detect_sections(self, text: str) -> List[Tuple[int, int, str]]:
        """
        Detect section boundaries in text.
        
        Args:
            text: Text to analyze
            
        Returns:
            List of (start_pos, end_pos, section_title) tuples
        """
        sections = []
        
        # Find all section headers
        for match in self.section_pattern.finditer(text):
            title = match.group(1) or match.group(2)
            start_pos = match.start()
            
            # Find the end of this section (start of next section or end of text)
            next_match = self.section_pattern.search(text, match.end())
            end_pos = next_match.start() if next_match else len(text)
            
            sections.append((start_pos, end_pos, title))
            
        return sections
    
    def _calculate_complexity(self, text: str) -> float:
        """
        Calculate text complexity to determine optimal chunk boundaries and overlap.
        
        Args:
            text: Text to analyze
            
        Returns:
            Complexity score (0.0-1.0)
        """
        if not text:
            return 0.0
            
        # Baseline complexity factors
        factors = []
        
        # 1. Sentence length variation
        sentences = self._tokenize_sentences(text)
        if sentences:
            sent_lengths = [len(s) for s in sentences]
            if len(sent_lengths) > 1:
                length_std = np.std(sent_lengths) / max(np.mean(sent_lengths), 1)
                factors.append(min(length_std, 1.0))
        
        # 2. Vocabulary richness
        words = re.findall(r'\b\w+\b', text.lower())
        if words:
            unique_ratio = len(set(words)) / len(words)
            factors.append(unique_ratio)
        
        # 3. Special content presence
        has_lists = bool(self.list_pattern.search(text))
        has_code = bool(self.code_block_pattern.search(text))
        has_tables = bool(self.table_pattern.search(text))
        
        special_content_score = sum([has_lists, has_code, has_tables]) / 3
        factors.append(special_content_score)
        
        # 4. Entity density
        if self.nlp:
            doc = self.nlp(text[:min(len(text), 5000)])  # Limit for performance
            entity_ratio = len(doc.ents) / max(len(doc), 1) * 5  # Scale up for impact
            factors.append(min(entity_ratio, 1.0))
        
        # Calculate overall complexity
        if factors:
            return sum(factors) / len(factors)
        return 0.5  # Default mid-complexity
    
    def _get_optimal_split_point(self, text: str, target_idx: int) -> int:
        """
        Find the optimal split point near the target index.
        
        Args:
            text: Text to split
            target_idx: Target split index
            
        Returns:
            Optimal split index
        """
        # Define the search window
        window_size = min(200, len(text) // 4)
        start_idx = max(0, target_idx - window_size)
        end_idx = min(len(text), target_idx + window_size)
        search_text = text[start_idx:end_idx]
        
        # Try to find paragraph boundaries first
        paragraphs = re.split(r'\n\s*\n', search_text)
        if len(paragraphs) > 1:
            # Find paragraph boundary closest to target
            current_pos = 0
            for para in paragraphs:
                if current_pos + len(para) + 2 > target_idx - start_idx:
                    return start_idx + current_pos
                current_pos += len(para) + 2
        
        # Next try sentence boundaries
        sentences = self._tokenize_sentences(search_text)
        if len(sentences) > 1:
            # Find sentence boundary closest to target
            current_pos = 0
            for sentence in sentences:
                next_pos = current_pos + len(sentence) + 1
                if next_pos >= target_idx - start_idx:
                    return start_idx + current_pos
                current_pos = next_pos
                
        # Fall back to word boundaries
        words = re.finditer(r'\b\w+\b', search_text)
        closest_word_boundary = target_idx
        for word in words:
            if word.start() <= target_idx - start_idx <= word.end():
                if abs(word.start() - (target_idx - start_idx)) < abs(word.end() - (target_idx - start_idx)):
                    closest_word_boundary = start_idx + word.start()
                else:
                    closest_word_boundary = start_idx + word.end()
                break
                
        return closest_word_boundary
    
    def _create_content_summary(self, text: str, max_length: int = 100) -> str:
        """
        Create a brief summary of content for chunk context.
        
        Args:
            text: Text to summarize
            max_length: Maximum summary length
            
        Returns:
            Brief summary
        """
        # Extract first sentence as a summary
        sentences = self._tokenize_sentences(text)
        if not sentences:
            return ""
            
        first_sentence = sentences[0]
        if len(first_sentence) <= max_length:
            return first_sentence
            
        # If first sentence is too long, truncate intelligently
        return first_sentence[:max_length-3] + "..."
        
    def _extract_key_entities(self, text: str, max_entities: int = 5) -> List[str]:
        """
        Extract key entities from text for context preservation.
        
        Args:
            text: Text to analyze
            max_entities: Maximum number of entities to extract
            
        Returns:
            List of key entity strings
        """
        entities = []
        
        if self.nlp:
            # Use spaCy for entity extraction
            doc = self.nlp(text[:min(len(text), 10000)])  # Limit for performance
            entity_counts = Counter([ent.text for ent in doc.ents])
            entities = [entity for entity, _ in entity_counts.most_common(max_entities)]
        else:
            # Simple fallback: extract capitalized phrases as potential entities
            capitalized = re.findall(r'\b[A-Z][a-zA-Z]+(?:\s+[A-Z][a-zA-Z]+)*\b', text)
            entities = [entity for entity, _ in Counter(capitalized).most_common(max_entities)]
            
        return entities
    
    def _detect_special_content(self, text: str) -> Dict[str, Any]:
        """
        Detect special content blocks for special handling.
        
        Args:
            text: Text to analyze
            
        Returns:
            Dictionary with special content markers
        """
        results = {
            "has_code": bool(self.code_block_pattern.search(text)),
            "has_tables": bool(self.table_pattern.search(text)),
            "has_lists": bool(self.list_pattern.search(text)),
            "special_content_ratio": 0.0
        }
        
        # Calculate ratio of special content
        special_content_length = 0
        
        for pattern in [self.code_block_pattern, self.table_pattern]:
            for match in pattern.finditer(text):
                special_content_length += match.end() - match.start()
                
        results["special_content_ratio"] = special_content_length / max(len(text), 1)
        
        return results
    
    async def chunk_text(self, text: str, metadata: Dict[str, Any]) -> List[Dict[str, str]]:
        """
        Chunk a text document using intelligent boundary detection and dynamic overlapping.
        
        Args:
            text: The text to chunk
            metadata: Metadata about the source document
            
        Returns:
            List of document chunks with enhanced metadata
        """
        # First pass: extract document structure
        sections = self._detect_sections(text) if self.respect_sections else []
        special_content = self._detect_special_content(text)
        
        # If document is small enough, return it as a single chunk
        if len(text) <= self.max_chunk_size:
            return [{
                "content": text,
                "metadata": {
                    **metadata,
                    "chunk_id": 0,
                    "chunk_strategy": "quantum",
                    "document_sections": [s[2] for s in sections],
                    "special_content": special_content,
                }
            }]
        
        # Prepare for chunking
        chunks = []
        current_pos = 0
        chunk_id = 0
        
        # Extract global document entities if relevant
        document_entities = self._extract_key_entities(text) if self.respect_entities else []
        
        # Process the text
        while current_pos < len(text):
            # Determine target end position for this chunk
            target_end = current_pos + self.chunk_size
            
            # Check if we're near the end of the document
            if target_end >= len(text):
                end_pos = len(text)
            else:
                # Find optimal split point
                end_pos = self._get_optimal_split_point(text, target_end)
            
            # Respect section boundaries if enabled
            if self.respect_sections and sections:
                # Check if we're splitting in the middle of a section
                for section_start, section_end, section_title in sections:
                    if current_pos < section_start < end_pos < section_end:
                        # We're splitting a section - adjust to section start if closer
                        if abs(section_start - end_pos) < abs(section_end - end_pos):
                            end_pos = section_start
                        else:
                            end_pos = section_end
                        break
            
            # Extract the chunk content
            chunk_text = text[current_pos:end_pos]
            
            # Calculate complexity for dynamic overlap
            complexity = self._calculate_complexity(chunk_text)
            
            # Determine actual overlap based on complexity
            if self.dynamic_overlap:
                # Scale overlap between 0.5x and 2x of base overlap based on complexity
                actual_overlap = int(self.chunk_overlap * (0.5 + 1.5 * complexity))
            else:
                actual_overlap = self.chunk_overlap
                
            # Identify current sections covered by this chunk
            chunk_sections = []
            for section_start, section_end, section_title in sections:
                if (current_pos <= section_start < end_pos) or (section_start <= current_pos < section_end):
                    chunk_sections.append(section_title)
            
            # Extract key entities in this chunk
            chunk_entities = self._extract_key_entities(chunk_text)
            
            # Create a content summary
            content_summary = self._create_content_summary(chunk_text)
            
            # Create enhanced metadata
            chunk_metadata = {
                **metadata,
                "chunk_id": chunk_id,
                "chunk_strategy": "quantum",
                "start_char": current_pos,
                "end_char": end_pos,
                "complexity": round(complexity, 2),
                "sections": chunk_sections,
                "summary": content_summary
            }
            
            # Add the entities if available
            if chunk_entities:
                chunk_metadata["entities"] = chunk_entities
                
            # Add document entities for context
            if document_entities:
                chunk_metadata["document_entities"] = document_entities
                
            # Include special content information
            if any(special_content.values()):
                for key, value in special_content.items():
                    if isinstance(value, bool) and value:
                        chunk_metadata[key] = value
            
            # Add the chunk
            chunks.append({
                "content": chunk_text,
                "metadata": chunk_metadata
            })
            
            # Update for next chunk with overlap
            current_pos = end_pos - actual_overlap if end_pos < len(text) else len(text)
            chunk_id += 1
            
            # Ensure we make progress
            if current_pos >= end_pos:
                current_pos = end_pos
                
        return chunks
        
    async def chunk_pdf(self, pdf_reader, file_path: str) -> List[Dict[str, str]]:
        """
        Chunk a PDF document with intelligent layout and structure analysis.
        
        Args:
            pdf_reader: The PDF reader object
            file_path: Path to the PDF file
            
        Returns:
            List of document chunks with enhanced metadata
        """
        # Extract text and structure from PDF
        all_text = ""
        page_texts = []
        page_breaks = []
        
        # Also track structural elements
        toc_items = []
        headers = []
        
        # Process each page
        for i, page in enumerate(pdf_reader.pages):
            page_text = page.extract_text()
            
            if page_text.strip():
                # Save the text and track page boundaries
                page_texts.append(page_text)
                all_text += page_text + "\n\n"
                page_breaks.append(len(all_text))
                
                # Try to detect headers by looking for larger fonts or bold text
                # This is simplified - real implementation would use PDF structure more
                try:
                    # Check for text with font size > 12 as potential headers
                    if hasattr(page, 'extract_text_lines'):
                        for line_obj in page.extract_text_lines():
                            if hasattr(line_obj, 'size') and line_obj.size > 12:
                                line_text = line_obj.get_text()
                                headers.append((len(all_text) - len(page_text) + page_text.find(line_text), 
                                              line_text))
                except:
                    # If structure extraction fails, use regex to guess headers
                    potential_headers = re.finditer(r'^[A-Z][^.!?]*(?:[.!?]|\n|$)', page_text, re.MULTILINE)
                    for match in potential_headers:
                        header_text = match.group(0).strip()
                        if len(header_text) < 100 and not header_text.endswith('.'):  # Avoid full sentences
                            position = len(all_text) - len(page_text) + match.start()
                            headers.append((position, header_text))
        
        # Attempt to extract TOC if available
        try:
            if hasattr(pdf_reader, 'outline') and pdf_reader.outline:
                for item in pdf_reader.outline:
                    if isinstance(item, dict) and 'title' in item and 'page' in item:
                        toc_items.append((item['title'], item['page']))
        except:
            pass  # TOC extraction is optional
        
        # Base metadata with enhanced PDF info
        base_metadata = {
            "source": Path(file_path).name,
            "document_type": "pdf",
            "total_pages": len(page_texts),
        }
        
        if toc_items:
            base_metadata["toc"] = toc_items
            
        # Use our advanced text chunking on the combined text
        chunks = await self.chunk_text(all_text, base_metadata)
        
        # Add page information to each chunk
        for chunk in chunks:
            chunk_start = chunk["metadata"]["start_char"]
            chunk_end = chunk["metadata"]["end_char"]
            
            # Find the page number(s)
            start_page = 1
            end_page = 1
            
            for i, break_pos in enumerate(page_breaks):
                if chunk_start < break_pos:
                    start_page = i + 1
                    break
                    
            for i, break_pos in enumerate(page_breaks):
                if chunk_end <= break_pos:
                    end_page = i + 1
                    break
            
            # Update metadata with page info
            chunk["metadata"]["page_start"] = start_page
            chunk["metadata"]["page_end"] = end_page
            
            # Add headers that appear in this chunk
            chunk_headers = []
            for header_pos, header_text in headers:
                if chunk_start <= header_pos < chunk_end:
                    chunk_headers.append(header_text)
                    
            if chunk_headers:
                chunk["metadata"]["headers"] = chunk_headers
            
        return chunks
        
    async def chunk_docx(self, doc, file_path: str) -> List[Dict[str, str]]:
        """
        Chunk a Word document using structure analysis and style detection.
        
        Args:
            doc: The Word document object
            file_path: Path to the Word file
            
        Returns:
            List of document chunks with enhanced metadata
        """
        # Extract text and structure
        all_text = ""
        paragraphs = []
        
        # Track structural elements
        headers = []
        bullet_points = []
        tables = []
        
        # Process each paragraph
        current_pos = 0
        for i, para in enumerate(doc.paragraphs):
            para_text = para.text.strip()
            if not para_text:
                continue
                
            paragraphs.append(para_text)
            
            # Track position
            para_start = current_pos
            all_text += para_text + "\n\n"
            current_pos = len(all_text)
            
            # Check for headers based on paragraph style
            if para.style.name.startswith('Heading'):
                headers.append((para_start, para_text, para.style.name))
                
            # Check for bullet points
            if para_text.startswith(('•', '-', '*')) or (para.style.name and 'List' in para.style.name):
                bullet_points.append((para_start, para_text))
                
        # Process tables if available
        table_idx = 0
        for table in doc.tables:
            table_text = ""
            for row in table.rows:
                for cell in row.cells:
                    table_text += cell.text + " | "
                table_text += "\n"
                
            if table_text.strip():
                tables.append(table_text)
                
                # Find approximate position in the document
                all_text += f"\n[TABLE {table_idx}]\n{table_text}\n\n"
                table_idx += 1
                
        # Base metadata with document structure info
        base_metadata = {
            "source": Path(file_path).name,
            "document_type": "docx",
            "has_tables": len(tables) > 0,
            "has_bullet_points": len(bullet_points) > 0,
            "has_headers": len(headers) > 0
        }
        
        # Use our advanced text chunking on the combined text
        chunks = await self.chunk_text(all_text, base_metadata)
        
        # Add structure information to each chunk
        for chunk in chunks:
            chunk_start = chunk["metadata"]["start_char"]
            chunk_end = chunk["metadata"]["end_char"]
            
            # Add headers that appear in this chunk
            chunk_headers = []
            for header_pos, header_text, header_style in headers:
                if chunk_start <= header_pos < chunk_end:
                    chunk_headers.append({
                        "text": header_text,
                        "style": header_style
                    })
                    
            if chunk_headers:
                chunk["metadata"]["headers"] = chunk_headers
                
            # Check for bullet points in this chunk
            has_bullets = any(chunk_start <= bp_pos < chunk_end for bp_pos, _ in bullet_points)
            if has_bullets:
                chunk["metadata"]["has_bullet_points"] = True
                
            # Estimate if a table is in this chunk
            if "[TABLE " in chunk["content"]:
                chunk["metadata"]["has_table"] = True
                
        return chunks
        
    async def chunk_tabular(self, df, file_path: str) -> List[Dict[str, str]]:
        """
        Chunk tabular data with intelligent column grouping and row segmentation.
        
        Args:
            df: The pandas DataFrame
            file_path: Path to the tabular file
            
        Returns:
            List of document chunks with enhanced metadata
        """
        chunks = []
        records = df.to_dict("records")
        
        # Skip if empty
        if not records:
            return []
            
        # Analyze dataframe structure
        column_types = {}
        numeric_columns = []
        text_columns = []
        date_columns = []
        
        for col in df.columns:
            if pd.api.types.is_numeric_dtype(df[col]):
                column_types[col] = "numeric"
                numeric_columns.append(col)
            elif pd.api.types.is_datetime64_dtype(df[col]):
                column_types[col] = "date"
                date_columns.append(col)
            else:
                column_types[col] = "text"
                text_columns.append(col)
                
        # Calculate column stats for context
        column_stats = {}
        for col in numeric_columns:
            column_stats[col] = {
                "min": float(df[col].min()) if not pd.isna(df[col].min()) else None,
                "max": float(df[col].max()) if not pd.isna(df[col].max()) else None,
                "mean": float(df[col].mean()) if not pd.isna(df[col].mean()) else None,
                "type": "numeric"
            }
            
        for col in date_columns:
            column_stats[col] = {
                "min": str(df[col].min()) if not pd.isna(df[col].min()) else None,
                "max": str(df[col].max()) if not pd.isna(df[col].max()) else None,
                "type": "date"
            }
            
        # Group rows based on a coherent strategy
        if len(records) <= self.max_chunk_size // 100:
            # Small dataframe - use a single chunk with all data
            content = df.to_string(index=False)
            chunks.append({
                "content": content,
                "metadata": {
                    "source": Path(file_path).name,
                    "document_type": "tabular",
                    "rows": f"1-{len(records)}",
                    "total_rows": len(records),
                    "columns": list(df.columns),
                    "column_types": column_types,
                    "column_stats": column_stats,
                    "chunk_id": 0,
                    "chunk_strategy": "quantum"
                }
            })
        else:
            # Strategy 1: Group by date if date column exists
            if date_columns and len(date_columns) > 0:
                # Sort by the first date column
                date_col = date_columns[0]
                sorted_df = df.sort_values(date_col)
                
                # Group by time periods
                try:
                    # Monthly groups
                    if 'year' in sorted_df[date_col].dt and 'month' in sorted_df[date_col].dt:
                        groups = sorted_df.groupby([sorted_df[date_col].dt.year, sorted_df[date_col].dt.month])
                        
                        for (year, month), group in groups:
                            group_records = group.to_dict("records")
                            if not group_records:
                                continue
                                
                            content = group.to_string(index=False)
                            chunks.append({
                                "content": content,
                                "metadata": {
                                    "source": Path(file_path).name,
                                    "document_type": "tabular",
                                    "row_range": f"{group.index.min()}-{group.index.max()}",
                                    "total_rows": len(group),
                                    "columns": list(group.columns),
                                    "year": year,
                                    "month": month,
                                    "column_types": column_types,
                                    "column_stats": column_stats,
                                    "chunk_id": len(chunks),
                                    "chunk_strategy": "quantum_date_grouped"
                                }
                            })
                except:
                    # Fallback to size-based chunking
                    self._chunk_tabular_by_size(df, file_path, column_types, column_stats, chunks)
            else:
                # No date column - use size-based chunking
                self._chunk_tabular_by_size(df, file_path, column_types, column_stats, chunks)
                
        return chunks
        
    def _chunk_tabular_by_size(self, df, file_path, column_types, column_stats, chunks):
        """
        Helper method to chunk tabular data by size.
        
        Args:
            df: The pandas DataFrame
            file_path: Path to the file
            column_types: Dictionary of column types
            column_stats: Dictionary of column statistics
            chunks: List to append chunks to
        """
        # Determine target rows per chunk
        target_rows = max(10, self.chunk_size // 100)  # Estimate 100 chars per row
        
        # Process in chunks
        current_chunk = []
        current_size = 0
        
        for i, row in enumerate(df.to_dict("records")):
            # Format the row as a string
            row_text = "\n".join([f"{k}: {v}" for k, v in row.items()])
            row_size = len(row_text)
            
            # If adding this row would exceed the chunk size and we already have content,
            # save the current chunk and start a new one
            if current_size + row_size > self.chunk_size and current_chunk:
                # Combine the rows into a single chunk
                content = "\n\n".join(current_chunk)
                chunks.append({
                    "content": content,
                    "metadata": {
                        "source": Path(file_path).name,
                        "document_type": "tabular",
                        "rows": f"{i-len(current_chunk)+1}-{i}",
                        "total_rows": len(current_chunk),
                        "columns": list(df.columns),
                        "column_types": column_types,
                        "column_stats": column_stats,
                        "chunk_id": len(chunks),
                        "chunk_strategy": "quantum_size_grouped"
                    }
                })
                
                # Start a new chunk
                current_chunk = []
                current_size = 0
                
            # Add the row to the current chunk
            current_chunk.append(row_text)
            current_size += row_size
                
        # Add the last chunk if it's not empty
        if current_chunk:
            content = "\n\n".join(current_chunk)
            chunks.append({
                "content": content,
                "metadata": {
                    "source": Path(file_path).name,
                    "document_type": "tabular",
                    "rows": f"{len(df)-len(current_chunk)+1}-{len(df)}",
                    "total_rows": len(current_chunk),
                    "columns": list(df.columns),
                    "column_types": column_types,
                    "column_stats": column_stats,
                    "chunk_id": len(chunks),
                    "chunk_strategy": "quantum_size_grouped"
                }
            })

class DocumentLoader:
    """
    Utility class for loading and processing documents from various file formats.
    """
    
    def __init__(self, upload_dir: str = "uploads", chunking_strategy: Literal["basic", "super", "quantum"] = "basic"):
        """
        Initialize the document loader.
        
        Args:
            upload_dir: Directory to store uploaded files
            chunking_strategy: Strategy to use for chunking documents
                - "basic": Simple chunking by paragraphs and pages
                - "super": Advanced semantic chunking with overlap
                - "quantum": Super Advanced semantic chunking
        """
        self.upload_dir = Path(upload_dir)
        self.upload_dir.mkdir(parents=True, exist_ok=True)
        
        # Set the chunking strategy
        self.set_chunking_strategy(chunking_strategy)
        
    def set_chunking_strategy(self, strategy: Literal["basic", "super", "quantum"]):
        """
        Set the chunking strategy.
        
        Args:
            strategy: The chunking strategy to use
        """
        if strategy == "super":
            self.chunker = SuperChunker()
        elif strategy == "quantum":
            self.chunker = QuantumChunker()
        else:
            self.chunker = BasicChunker()
        
    async def save_uploaded_file(self, file) -> str:
        """
        Save an uploaded file to disk and return its path.
        
        Args:
            file: The uploaded file object from FastAPI
            
        Returns:
            Path to the saved file
        """
        # Create a unique filename to avoid collisions
        filename = f"{uuid.uuid4()}_{file.filename}"
        file_path = self.upload_dir / filename
        
        # Write the file to disk
        with open(file_path, "wb") as f:
            content = await file.read()
            f.write(content)
            
        return str(file_path)
    
    async def load_document(self, file_path: str) -> List[Dict[str, str]]:
        """
        Load a document and split it into chunks.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            List of document chunks with metadata
        """
        file_ext = Path(file_path).suffix.lower()
        
        try:
            if file_ext == ".txt":
                return await self._load_text(file_path)
            elif file_ext == ".pdf":
                return await self._load_pdf(file_path)
            elif file_ext in [".docx", ".doc"]:
                return await self._load_docx(file_path)
            elif file_ext in [".csv", ".xlsx", ".xls"]:
                return await self._load_tabular(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_ext}")
        except Exception as e:
            print(f"Error loading document {file_path}: {str(e)}")
            return []
            
    async def _load_text(self, file_path: str) -> List[Dict[str, str]]:
        """
        Load and chunk a text file.
        """
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
            
        # Use the chunking strategy
        metadata = {"source": Path(file_path).name}
        return await self.chunker.chunk_text(text, metadata)
        
    async def _load_pdf(self, file_path: str) -> List[Dict[str, str]]:
        """
        Load and chunk a PDF file.
        """
        try:
            from pypdf import PdfReader
            
            reader = PdfReader(file_path)
            return await self.chunker.chunk_pdf(reader, file_path)
            
        except ImportError:
            raise ImportError("pypdf is required for PDF processing. Install with: pip install pypdf")
            
    async def _load_docx(self, file_path: str) -> List[Dict[str, str]]:
        """
        Load and chunk a Word document.
        """
        try:
            import docx
            
            doc = docx.Document(file_path)
            return await self.chunker.chunk_docx(doc, file_path)
            
        except ImportError:
            raise ImportError("python-docx is required for DOCX processing. Install with: pip install python-docx")
            
    async def _load_tabular(self, file_path: str) -> List[Dict[str, str]]:
        """
        Load and process tabular data (CSV, Excel).
        """
        file_ext = Path(file_path).suffix.lower()
        
        try:
            import pandas as pd
            
            if file_ext == ".csv":
                df = pd.read_csv(file_path)
            else:  # Excel files
                df = pd.read_excel(file_path)
                
            return await self.chunker.chunk_tabular(df, file_path)
            
        except ImportError:
            raise ImportError("pandas is required for tabular data processing. Install with: pip install pandas openpyxl")